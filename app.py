# ✅ STEP 1: Cài đặt môi trường Python

# Tạo môi trường ảo (tuỳ chọn, nhưng nên dùng)
# python -m venv venv
# venv\Scripts\activate (Windows) hoặc source venv/bin/activate (Mac/Linux)

# Cài thư viện cần thiết:
# pip install streamlit PyMuPDF deep-translator python-docx

# ✅ STEP 2: Tạo file app.py

import streamlit as st
import fitz  # PyMuPDF
from deep_translator import GoogleTranslator
from docx import Document
from io import BytesIO

st.set_page_config(page_title="PDF Translator (EN→VI)", layout="wide")
st.title("📘 Dịch tài liệu PDF từ tiếng Anh sang tiếng Việt")

# Khởi tạo trạng thái trang
if "page_number" not in st.session_state:
    st.session_state.page_number = 1

if "translated_pages" not in st.session_state:
    st.session_state.translated_pages = {}

uploaded_file = st.file_uploader("Tải lên file PDF tiếng Anh:", type="pdf")

if uploaded_file:
    pdf_doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
    total_pages = len(pdf_doc)

    # Xử lý nút Prev / Next
    col_nav = st.columns([1, 1])
    with col_nav[0]:
        if st.button("⬅️ Prev page"):
            if st.session_state.page_number > 1:
                st.session_state.page_number -= 1
    with col_nav[1]:
        if st.button("Next page ➡️"):
            if st.session_state.page_number < total_pages:
                st.session_state.page_number += 1

    # Lấy nội dung trang hiện tại
    page_number = st.session_state.page_number
    page = pdf_doc[page_number - 1]
    html_text = page.get_text("html")
    text_plain = page.get_text()

    col1, col2 = st.columns(2)

    # Style CSS cho box
    box_style = """
        border: 1px solid #a3c4f3;
        border-radius: 10px;
        padding: 10px;
        background-color: #f9fbff;
        font-size: 10pt;
        line-height: 1.6;
        overflow-x: auto;
    """

    with col1:
        st.markdown(f"**📄 Trang {page_number} / {total_pages}**")
        st.markdown(f"<div style='{box_style}'>{html_text}</div>", unsafe_allow_html=True)

    with col2:
        st.markdown(f"**🌐 Bản dịch tiếng Việt (Trang {page_number}):**")
        if text_plain.strip() != "":
            with st.spinner("⏳ Đang dịch bằng Google Translate..."):
                try:
                    translated = GoogleTranslator(source='en', target='vi').translate(text_plain)
                    translated_html = translated.replace("\n", "<br>")
                    st.markdown(f"<div style='{box_style}'>{translated_html}</div>", unsafe_allow_html=True)
                    st.session_state.translated_pages[page_number] = translated
                except Exception as e:
                    st.error(f"Lỗi khi dịch: {e}")
        else:
            st.warning("Không có nội dung để dịch.")

    # Nút Export .docx
    st.markdown("<div style='text-align:right; padding-top:10px;'>", unsafe_allow_html=True)
    if st.button("📥 Export .docx"):
        doc = Document()
        doc.add_heading("Bản dịch PDF sang tiếng Việt", level=1)
        for i in range(1, total_pages + 1):
            if i in st.session_state.translated_pages:
                doc.add_page_break()
                doc.add_heading(f"Trang {i}", level=2)
                doc.add_paragraph(st.session_state.translated_pages[i])
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        st.download_button(
            label="📄 Tải file .docx",
            data=buffer,
            file_name="ban_dich.pdf.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    st.markdown("</div>", unsafe_allow_html=True)

# ✅ STEP 3: Chạy ứng dụng
# streamlit run app.py

# ✅ STEP 4: Đưa lên GitHub
# - Tạo repo GitHub
# - Thêm file app.py
# - Tạo file requirements.txt chứa:
# streamlit
# PyMuPDF
# deep-translator
# python-docx

# ✅ Bổ sung file .gitignore:
# __pycache__/
# *.pyc
# venv/
# .streamlit/secrets.toml
